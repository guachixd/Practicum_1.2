"""
extractor.py
------------
Extrae texto plano de los documentos fuente de la base de conocimiento
(plan docente, guía didáctica, materiales de unidad) para que puedan
ser consumidos por base_conocimiento.py.

Formatos soportados:
    - PDF   -> pypdf
    - DOCX  -> python-docx (por si la guía didáctica viene en Word)

Un documento puede tener páginas sin texto extraíble (diapositivas que
son solo imágenes, escaneos, etc.). Ese contenido se descarta y se
reporta con una advertencia, porque procesarlo requeriría OCR
(pytesseract u otra herramienta) que está fuera del alcance de este
pipeline.
"""

from __future__ import annotations

import os
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument


# Un umbral bajo de caracteres para considerar que una página "no tiene
# texto útil" (evita quedarnos con basura de 2-3 caracteres sueltos).
MIN_CARACTERES_UTILES = 20


def extraer_texto_pdf(ruta_pdf: str) -> str:
    """
    Extrae el texto de todas las páginas de un PDF y lo concatena.

    Si una página no tiene texto extraíble (por ejemplo, una diapositiva
    que es solo una imagen), se omite y se informa por consola en lugar
    de insertar contenido vacío o inventado.
    """
    ruta = Path(ruta_pdf)
    if not ruta.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {ruta_pdf}")

    lector = PdfReader(str(ruta))
    paginas_con_texto = []
    paginas_descartadas = 0

    for i, pagina in enumerate(lector.pages, start=1):
        texto = (pagina.extract_text() or "").strip()
        if len(texto) >= MIN_CARACTERES_UTILES:
            paginas_con_texto.append(texto)
        else:
            paginas_descartadas += 1

    if paginas_descartadas:
        print(
            f"[extractor] Aviso: {paginas_descartadas} página(s) de "
            f"'{ruta.name}' no tenían texto extraíble y fueron "
            f"descartadas (posibles imágenes/diapositivas escaneadas)."
        )

    return "\n\n".join(paginas_con_texto)


def extraer_texto_pdf_por_pagina(ruta_pdf: str) -> list[str]:
    """
    Igual que extraer_texto_pdf, pero devuelve una lista con el texto de
    cada página (útil cuando luego se necesita segmentar por rangos de
    páginas, no solo por palabras clave).
    """
    ruta = Path(ruta_pdf)
    if not ruta.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {ruta_pdf}")

    lector = PdfReader(str(ruta))
    return [(pagina.extract_text() or "").strip() for pagina in lector.pages]


def extraer_texto_docx(ruta_docx: str) -> str:
    """
    Extrae el texto de un documento Word (párrafos y celdas de tablas).
    """
    ruta = Path(ruta_docx)
    if not ruta.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {ruta_docx}")

    doc = DocxDocument(str(ruta))
    partes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    partes.append(celda.text.strip())

    return "\n".join(partes)


def extraer_documento(ruta: str) -> str:
    """
    Punto de entrada único: detecta la extensión del archivo y llama al
    extractor correspondiente.
    """
    extension = Path(ruta).suffix.lower()

    if extension == ".pdf":
        return extraer_texto_pdf(ruta)
    if extension in (".docx", ".doc"):
        return extraer_texto_docx(ruta)

    raise ValueError(
        f"Formato no soportado: '{extension}'. "
        "Este pipeline solo procesa PDF y Word (sin OCR)."
    )


def extraer_carpeta_materiales(ruta_carpeta: str) -> dict[str, str]:
    """
    Extrae el texto de todos los documentos dentro de data/materiales/.

    Devuelve un diccionario {nombre_de_archivo: texto_extraido}.
    Los archivos que no puedan procesarse (formato no soportado, PDF sin
    texto extraíble, etc.) se omiten con una advertencia en vez de
    detener el pipeline completo.
    """
    carpeta = Path(ruta_carpeta)
    resultado: dict[str, str] = {}

    if not carpeta.exists():
        print(f"[extractor] Aviso: la carpeta {ruta_carpeta} no existe.")
        return resultado

    for archivo in sorted(carpeta.iterdir()):
        if archivo.is_dir():
            continue
        try:
            texto = extraer_documento(str(archivo))
        except ValueError as e:
            print(f"[extractor] Aviso: se omite '{archivo.name}' ({e})")
            continue

        if len(texto.strip()) < MIN_CARACTERES_UTILES:
            print(
                f"[extractor] Aviso: '{archivo.name}' no aportó texto "
                "útil (probablemente son solo imágenes/diapositivas "
                "escaneadas) y se descarta de la base de conocimiento."
            )
            continue

        resultado[archivo.stem] = texto

    return resultado


if __name__ == "__main__":
    # Prueba rápida manual: python src/extractor.py
    base_dir = Path(__file__).resolve().parent.parent
    texto_plan = extraer_texto_pdf(str(base_dir / "data" / "plan_docente.pdf"))
    print(f"Plan docente: {len(texto_plan)} caracteres extraídos.")

    materiales = extraer_carpeta_materiales(str(base_dir / "data" / "materiales"))
    for nombre, texto in materiales.items():
        print(f"Material '{nombre}': {len(texto)} caracteres extraídos.")
