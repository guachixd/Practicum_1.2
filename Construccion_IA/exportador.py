"""
exportador.py — Exportador

Corresponde al contenedor "Exportador" de tu diagrama C4.
Lee los recursos generados (desde MongoDB, vía bd.py) y los ensambla en un
único documento Word: portada + introducción general + 8 unidades + subsecciones.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bd import conectar_bd, obtener_recursos_unidad, obtener_introduccion

CARPETA_SALIDA = os.path.join(os.path.dirname(__file__), "..", "data", "salida")


def _agregar_portada(doc: Document, asignatura: str, codigo: str):
    doc.add_paragraph().add_run().add_break()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Recursos Didácticos")
    run.bold = True
    run.font.size = Pt(28)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run(f"{asignatura} ({codigo})")
    run2.font.size = Pt(16)

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nota.add_run("Generado automáticamente a partir del material del curso.").italic = True

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(date.today().strftime("%d/%m/%Y")).italic = True

    doc.add_page_break()


def _agregar_introduccion(doc: Document, introduccion: dict):
    doc.add_heading("Introducción a la asignatura", level=1)
    texto = introduccion.get("introduccion", "")
    if texto:
        for parrafo in texto.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
    else:
        doc.add_paragraph("(No disponible)")
    doc.add_page_break()


def _agregar_unidad(doc: Document, numero: int, titulo_unidad: str, recursos: dict):
    doc.add_heading(f"Unidad {numero}: {titulo_unidad}", level=1)

    if not recursos:
        doc.add_paragraph("No se generaron recursos para esta unidad (sin material disponible).")
        doc.add_page_break()
        return

    # --- Resumen ejecutivo ---
    doc.add_heading("Resumen ejecutivo", level=2)
    resumen = recursos.get("resumen", {})
    parrafo_1 = resumen.get("parrafo_1", "")
    parrafo_2 = resumen.get("parrafo_2", "")
    if parrafo_1 or parrafo_2:
        if parrafo_1:
            doc.add_paragraph(parrafo_1)
        if parrafo_2:
            doc.add_paragraph(parrafo_2)
    else:
        doc.add_paragraph("(No disponible)")

    # --- Glosario ---
    doc.add_heading("Glosario", level=2)
    glosario = recursos.get("glosario", {})
    terminos = glosario.get("terminos", [])
    if terminos:
        for t in terminos:
            p = doc.add_paragraph()
            p.add_run(f"{t.get('termino', '')}: ").bold = True
            p.add_run(t.get("definicion", ""))
    else:
        doc.add_paragraph("(No disponible)")

    # --- Preguntas de comprensión ---
    doc.add_heading("Preguntas de comprensión", level=2)
    preguntas_data = recursos.get("preguntas", {})
    preguntas = preguntas_data.get("preguntas", [])
    etiquetas_nivel = {"facil": "Fácil", "media": "Media", "dificil": "Difícil"}
    if preguntas:
        for i, q in enumerate(preguntas, start=1):
            nivel = etiquetas_nivel.get(q.get("nivel", ""), q.get("nivel", ""))
            p = doc.add_paragraph()
            etiqueta = f" [{nivel}]" if nivel else ""
            p.add_run(f"{i}. {q.get('pregunta', '')}{etiqueta}").bold = True
            doc.add_paragraph(f"   Respuesta esperada: {q.get('respuesta_esperada', '')}")
    else:
        doc.add_paragraph("(No disponible)")

    # --- Actividad de reflexión ---
    doc.add_heading("Actividad de reflexión", level=2)
    actividad = recursos.get("actividad", {})
    if actividad.get("consigna"):
        doc.add_paragraph(actividad["consigna"])
        p = doc.add_paragraph()
        p.add_run("Entregable esperado: ").bold = True
        p.add_run(actividad.get("entregable_esperado", ""))
    else:
        doc.add_paragraph("(No disponible)")

    doc.add_page_break()


def exportar_documento_final(unidades: list[dict], asignatura: str, codigo: str,
                              ruta_salida: str = None):
    """
    Arma el documento Word final leyendo los recursos ya generados de MongoDB
    para la introducción general y cada una de las 8 unidades.
    """
    if ruta_salida is None:
        os.makedirs(CARPETA_SALIDA, exist_ok=True)
        ruta_salida = os.path.join(CARPETA_SALIDA, "COMP2010_recursos_didacticos.docx")

    db = conectar_bd()
    if db is None:
        raise RuntimeError("No se pudo conectar a la base de datos. Verifica que MongoDB esté corriendo.")

    doc = Document()
    _agregar_portada(doc, asignatura, codigo)

    introduccion = obtener_introduccion(db)
    _agregar_introduccion(doc, introduccion)

    for unidad in unidades:
        recursos = obtener_recursos_unidad(db, unidad["numero"])
        if not recursos:
            print(f"  Unidad {unidad['numero']}: sin recursos generados, se omite del documento final.")
            continue
        _agregar_unidad(doc, unidad["numero"], unidad["titulo"], recursos)

    doc.save(ruta_salida)
    print(f"Documento final generado: {os.path.abspath(ruta_salida)}")


# ==========================================
# BLOQUE DE PRUEBA LOCAL (con datos simulados, sin Gemini ni Mongo real)
# ==========================================
if __name__ == "__main__":
    doc = Document()
    _agregar_portada(doc, "Arquitectura y Organización de Computadores", "COMP_2010")
    _agregar_introduccion(doc, {"introduccion": "Esta asignatura cubre los fundamentos de arquitectura y organización de computadores."})

    recursos_ejemplo = {
        "resumen": {
            "parrafo_1": "La arquitectura define lo visible al programador, como el conjunto de instrucciones.",
            "parrafo_2": "La organización define la implementación física: unidades operativas y su interconexión."
        },
        "glosario": {"terminos": [
            {"termino": "Arquitectura", "definicion": "Atributos visibles al programador."},
            {"termino": "Organización", "definicion": "Unidades operativas y su interconexión."}
        ]},
        "preguntas": {"preguntas": [
            {"pregunta": "¿Qué es la arquitectura de un computador?",
             "respuesta_esperada": "Los atributos visibles al programador.",
             "nivel": "facil"},
            {"pregunta": "¿Cómo se relacionan arquitectura y organización en un mismo procesador?",
             "respuesta_esperada": "La arquitectura se mantiene estable mientras la organización evoluciona con la tecnología.",
             "nivel": "media"},
            {"pregunta": "¿Podría cambiar la organización de un computador sin alterar su arquitectura? Justifica.",
             "respuesta_esperada": "Sí, porque la organización es la implementación física y puede optimizarse sin cambiar lo visible al programador.",
             "nivel": "dificil"}
        ]},
        "actividad": {
            "consigna": "Compara la arquitectura x86 de dos procesadores distintos.",
            "entregable_esperado": "Un cuadro comparativo de una página."
        }
    }

    _agregar_unidad(doc, 1, "Organización vs Arquitectura", recursos_ejemplo)
    _agregar_unidad(doc, 2, "Unidad sin material", {})

    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    ruta_prueba = os.path.join(CARPETA_SALIDA, "prueba_exportador.docx")
    doc.save(ruta_prueba)
    print(f"Documento de prueba generado: {ruta_prueba}")
