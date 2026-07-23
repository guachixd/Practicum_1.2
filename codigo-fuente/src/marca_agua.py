"""
marca_agua.py — Marca de agua institucional

Protección de propiedad intelectual: cada documento Word generado lleva
una marca de agua con el logo de la UTPL y el nombre del docente, visible
de fondo en TODAS las páginas del documento (va en el encabezado, que se
repite automáticamente).

Estrategia:
1. Se compone una sola imagen PNG con transparencia: el logo de la UTPL
   desvanecido + una línea de texto con el nombre del docente, rotada en
   diagonal (como cualquier marca de agua estándar de Word).
2. Esa imagen se inserta en el encabezado del documento como una imagen
   flotante "detrás del texto" (behindDoc), centrada en la página — la
   misma técnica que usa Word internamente para sus marcas de agua.
   python-docx no expone esto en su API pública, así que se arma el XML
   (WordprocessingML/DrawingML) directamente sobre el run ya insertado.
"""

import os
from PIL import Image, ImageDraw, ImageFont

from docx.shared import Emu, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OPACIDAD_LOGO = 40       # 0-255. Bajo = más tenue (no debe estorbar la lectura).
OPACIDAD_TEXTO = 110      # el texto puede ir un poco más marcado que el logo
ANCHO_LOGO_PX = 520
ANGULO_ROTACION = 35     # grados, look clásico de marca de agua diagonal


def _cargar_fuente(tamano: int) -> ImageFont.FreeTypeFont:
    """Busca una fuente TrueType razonable en el sistema; si no encuentra
    ninguna, cae al font por defecto de Pillow (más feo, pero no rompe)."""
    candidatas = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
    ]
    for ruta in candidatas:
        if os.path.isfile(ruta):
            return ImageFont.truetype(ruta, tamano)
    return ImageFont.load_default()


def generar_imagen_marca_agua(ruta_logo: str, nombre_docente: str, ruta_salida: str) -> str:
    """
    Compone la imagen de marca de agua (logo UTPL desvanecido + nombre del
    docente) ya rotada, lista para insertarse en el encabezado del Word.
    Devuelve la ruta del PNG generado.
    """
    logo = Image.open(ruta_logo).convert("RGBA")
    ratio = ANCHO_LOGO_PX / logo.width
    logo = logo.resize((ANCHO_LOGO_PX, int(logo.height * ratio)), Image.LANCZOS)

    # Desvanece el logo (reduce el canal alfa proporcionalmente).
    r, g, b, a = logo.split()
    a = a.point(lambda px: int(px * (OPACIDAD_LOGO / 255)))
    logo = Image.merge("RGBA", (r, g, b, a))

    texto_linea1 = "UNIVERSIDAD TÉCNICA PARTICULAR DE LOJA"
    texto_linea2 = f"Elaborado por: {nombre_docente}"
    texto_linea3 = "Documento de uso académico — protegido por derechos de autor"

    fuente_1 = _cargar_fuente(26)
    fuente_2 = _cargar_fuente(22)
    fuente_3 = _cargar_fuente(16)

    # Lienzo de trabajo (sin rotar todavía): logo arriba, texto debajo.
    ancho_lienzo = max(ANCHO_LOGO_PX, 700)
    alto_lienzo = logo.height + 120
    lienzo = Image.new("RGBA", (ancho_lienzo, alto_lienzo), (255, 255, 255, 0))

    lienzo.paste(logo, ((ancho_lienzo - logo.width) // 2, 0), logo)

    dibujo = ImageDraw.Draw(lienzo)
    color_texto = (10, 46, 82, OPACIDAD_TEXTO)  # azul institucional UTPL, desvanecido

    def _centrar_y_dibujar(texto, fuente, y):
        ancho_texto = dibujo.textlength(texto, font=fuente)
        x = (ancho_lienzo - ancho_texto) / 2
        dibujo.text((x, y), texto, font=fuente, fill=color_texto)

    _centrar_y_dibujar(texto_linea1, fuente_1, logo.height + 8)
    _centrar_y_dibujar(texto_linea2, fuente_2, logo.height + 42)
    _centrar_y_dibujar(texto_linea3, fuente_3, logo.height + 72)

    # Rota todo el conjunto (logo + texto) en diagonal, expandiendo el
    # lienzo para no recortar esquinas.
    lienzo_rotado = lienzo.rotate(ANGULO_ROTACION, expand=True, resample=Image.BICUBIC)

    os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
    lienzo_rotado.save(ruta_salida, "PNG")
    return ruta_salida


def _convertir_a_flotante_detras_del_texto(run):
    """
    Toma el run que ya tiene una imagen insertada (inline, vía
    run.add_picture) y reescribe su XML para que la imagen quede como
    objeto flotante, centrado en la página y DETRÁS del texto — el
    comportamiento estándar de una marca de agua en Word. python-docx no
    tiene una API pública para esto, así que se manipula el árbol
    WordprocessingML directamente.
    """
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    cnv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "1")
    anchor.set("behindDoc", "1")     # <- clave: la imagen queda detrás del texto
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    align_h = OxmlElement("wp:align")
    align_h.text = "center"
    position_h.append(align_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    align_v = OxmlElement("wp:align")
    align_v.text = "center"
    position_v.append(align_v)

    wrap_none = OxmlElement("wp:wrapNone")

    anchor.append(simple_pos)
    anchor.append(position_h)
    anchor.append(position_v)
    anchor.append(extent)
    anchor.append(wrap_none)
    # El esquema de wp:anchor exige este orden exacto: docPr ANTES de
    # cNvGraphicFramePr. Invertirlo (como estaba antes) genera un .docx que
    # Word considera corrupto y se niega a abrir ("Word experienced an
    # error trying to open the file"), aunque el archivo se haya guardado
    # sin ningún error de Python.
    anchor.append(doc_pr)
    if cnv_graphic_frame_pr is not None:
        anchor.append(cnv_graphic_frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def agregar_marca_agua(doc, ruta_logo: str, nombre_docente: str, carpeta_temporal: str):
    """
    Genera la imagen de marca de agua y la inserta en el encabezado de
    CADA sección del documento (por eso aparece en todas las páginas,
    incluida la portada).
    """
    ruta_imagen = os.path.join(carpeta_temporal, "_marca_agua.png")
    generar_imagen_marca_agua(ruta_logo, nombre_docente, ruta_imagen)

    for seccion in doc.sections:
        seccion.header.is_linked_to_previous = False
        encabezado = seccion.header
        parrafo = encabezado.paragraphs[0] if encabezado.paragraphs else encabezado.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = parrafo.add_run()
        run.add_picture(ruta_imagen, width=Inches(5.2))
        _convertir_a_flotante_detras_del_texto(run)
