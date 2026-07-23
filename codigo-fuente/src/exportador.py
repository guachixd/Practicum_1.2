"""
exportador.py — Exportador

Corresponde al contenedor "Exportador" de tu diagrama C4.
Lee los recursos generados (desde MongoDB, vía bd.py) y los ensambla en un
único documento Word: portada + introducción general + unidades del curso
(cuantas hayan salido del plan docente, ya no fijas en 8) + marca de agua
institucional en cada página.
"""

import os
import tempfile
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bd import obtener_recursos_unidad, obtener_introduccion
from marca_agua import agregar_marca_agua

CARPETA_SALIDA = os.path.join(os.path.dirname(__file__), "..", "data", "salida")

AZUL_UTPL = RGBColor(0x0A, 0x2A, 0x45)


def _sombrear_celda(celda, color_hex: str):
    """Le pone color de fondo a una celda de tabla (python-docx no trae esto de fábrica)."""
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(sombra)


def _agregar_portada(doc: Document, asignatura: str, codigo: str, nombre_docente: str,
                      ruta_logo: str = None, carrera: str = "", ciclo_academico: str = "", periodo: str = ""):
    doc.add_paragraph().add_run().add_break()

    if ruta_logo and os.path.isfile(ruta_logo):
        parrafo_logo = doc.add_paragraph()
        parrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo_logo.add_run().add_picture(ruta_logo, width=Inches(2.6))
        doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Recursos Didácticos")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = AZUL_UTPL

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run(f"{asignatura} ({codigo})")
    run2.font.size = Pt(20)
    run2.bold = True

    # Carrera, ciclo y periodo van cada uno en su propia línea, grandes,
    # porque son justo el tipo de dato que un docente quiere ver de un
    # vistazo en la portada (para qué grupo y periodo es este documento).
    for texto in (carrera, f"Ciclo {ciclo_academico}" if ciclo_academico else "", periodo):
        if texto:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(texto).font.size = Pt(14)

    docente = doc.add_paragraph()
    docente.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_docente = docente.add_run(f"Elaborado por: {nombre_docente}")
    run_docente.font.size = Pt(14)
    run_docente.bold = True

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nota.add_run("Generado automáticamente a partir del material del curso.").italic = True

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(date.today().strftime("%d/%m/%Y")).italic = True

    doc.add_page_break()


def _agregar_introduccion(doc: Document, introduccion: dict):
    doc.add_heading("Introducción a la asignatura", level=1)
    texto = introduccion.get("introduccion", "")
    if texto:
        for parrafo in texto.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
    else:
        doc.add_paragraph("(No disponible)")
    doc.add_page_break()


def _agregar_unidad(doc: Document, numero: int, titulo_unidad: str, recursos: dict,
                     tiene_material: bool = True):
    doc.add_heading(f"Unidad {numero}: {titulo_unidad}", level=1)

    if not tiene_material:
        aviso = doc.add_paragraph()
        run_aviso = aviso.add_run(
            "Nota: el docente no subió material propio para esta unidad. El contenido a "
            "continuación se generó a partir del plan docente (título y temas de la unidad), "
            "no fue verificado contra material fuente específico."
        )
        run_aviso.italic = True
        run_aviso.font.size = Pt(9)

    if not recursos:
        doc.add_paragraph("No se generaron recursos para esta unidad.")
        doc.add_page_break()
        return

    # --- Resumen ejecutivo ---
    doc.add_heading("Resumen ejecutivo", level=2)
    resumen = recursos.get("resumen", {})
    parrafo_1 = resumen.get("parrafo_1", "")
    parrafo_2 = resumen.get("parrafo_2", "")
    if parrafo_1 or parrafo_2:
        if parrafo_1:
            doc.add_paragraph(parrafo_1)
        if parrafo_2:
            doc.add_paragraph(parrafo_2)
    else:
        doc.add_paragraph("(No disponible)")

    # --- Glosario (como tabla, no como párrafos sueltos) ---
    doc.add_heading("Glosario", level=2)
    glosario = recursos.get("glosario", {})
    terminos = glosario.get("terminos", [])
    if terminos:
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = "Table Grid"
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        encabezado = tabla.rows[0].cells
        encabezado[0].text = "Término"
        encabezado[1].text = "Definición"
        for celda in encabezado:
            _sombrear_celda(celda, "0A2A45")
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True

        for t in terminos:
            fila = tabla.add_row().cells
            fila[0].text = t.get("termino", "")
            fila[0].paragraphs[0].runs[0].bold = True
            fila[1].text = t.get("definicion", "")
    else:
        doc.add_paragraph("(No disponible)")
    doc.add_paragraph()

    # --- Preguntas de comprensión ---
    doc.add_heading("Preguntas de comprensión", level=2)
    preguntas_data = recursos.get("preguntas", {})
    preguntas = preguntas_data.get("preguntas", [])
    etiquetas_nivel = {"facil": "Fácil", "media": "Media", "dificil": "Difícil"}
    if preguntas:
        for i, q in enumerate(preguntas, start=1):
            nivel = etiquetas_nivel.get(q.get("nivel", ""), q.get("nivel", ""))
            p = doc.add_paragraph()
            etiqueta = f" [{nivel}]" if nivel else ""
            p.add_run(f"{i}. {q.get('pregunta', '')}{etiqueta}").bold = True
            doc.add_paragraph(f"   Respuesta esperada: {q.get('respuesta_esperada', '')}")
    else:
        doc.add_paragraph("(No disponible)")

    # --- Actividad de reflexión ---
    doc.add_heading("Actividad de reflexión", level=2)
    actividad = recursos.get("actividad", {})
    if actividad.get("consigna"):
        doc.add_paragraph(actividad["consigna"])
        p = doc.add_paragraph()
        p.add_run("Entregable esperado: ").bold = True
        p.add_run(actividad.get("entregable_esperado", ""))
    else:
        doc.add_paragraph("(No disponible)")

    doc.add_page_break()


def exportar_documento_final(db, job_id: str, unidades: list[dict], asignatura: str, codigo: str,
                              nombre_docente: str = "Docente no especificado",
                              ruta_logo: str = None,
                              ruta_salida: str = None,
                              carrera: str = "", ciclo_academico: str = "", periodo: str = ""):
    """
    Arma el documento Word final leyendo los recursos ya generados de
    MongoDB (aislados por job_id) para la introducción general y cada una
    de las unidades del curso (la cantidad ya no es fija: sale del plan
    docente).

    Recibe la conexión `db` ya abierta (en vez de crear una nueva) para
    reutilizar la misma conexión que usó el resto del pipeline.

    El logo se usa en dos lugares distintos: siempre en grande en la
    portada (si se indica una ruta válida), y además como marca de agua
    tenue en el encabezado de cada página, pero SOLO si el docente eligió
    esa opción antes de generar (ruta_logo llega como None desde el
    pipeline cuando el docente dijo que no quería marca de agua).
    """
    if ruta_salida is None:
        os.makedirs(CARPETA_SALIDA, exist_ok=True)
        ruta_salida = os.path.join(CARPETA_SALIDA, f"{codigo}_{job_id}_recursos_didacticos.docx")
    else:
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)

    ruta_logo_portada = ruta_logo
    if not ruta_logo_portada:
        ruta_logo_portada = os.path.join(os.path.dirname(__file__), "..", "assets", "logo_utpl.png")
        if not os.path.isfile(ruta_logo_portada):
            ruta_logo_portada = None

    doc = Document()
    _agregar_portada(doc, asignatura, codigo, nombre_docente, ruta_logo_portada, carrera, ciclo_academico, periodo)

    introduccion = obtener_introduccion(db, job_id)
    _agregar_introduccion(doc, introduccion)

    for unidad in unidades:
        recursos = obtener_recursos_unidad(db, job_id, unidad["numero"])
        if not recursos:
            print(f"  Unidad {unidad['numero']}: sin recursos generados, se omite del documento final.")
            continue
        _agregar_unidad(doc, unidad["numero"], unidad["titulo"], recursos, unidad["tiene_material"])

    if ruta_logo and os.path.isfile(ruta_logo):
        with tempfile.TemporaryDirectory() as carpeta_temporal:
            agregar_marca_agua(doc, ruta_logo, nombre_docente, carpeta_temporal)
            doc.save(ruta_salida)
    else:
        print("  [!] No se indicó logo institucional (o no se encontró); el documento se genera sin marca de agua.")
        doc.save(ruta_salida)

    print(f"Documento final generado: {os.path.abspath(ruta_salida)}")
    return ruta_salida


# ==========================================
# BLOQUE DE PRUEBA LOCAL (con datos simulados, sin Gemini ni Mongo real)
# ==========================================
if __name__ == "__main__":
    doc = Document()
    _agregar_portada(doc, "Arquitectura y Organización de Computadores", "COMP_2010", "Ing. Juan Pérez Rodríguez")
    _agregar_introduccion(doc, {"introduccion": "Esta asignatura cubre los fundamentos de arquitectura y organización de computadores."})

    recursos_ejemplo = {
        "resumen": {
            "parrafo_1": "La arquitectura define lo visible al programador, como el conjunto de instrucciones.",
            "parrafo_2": "La organización define la implementación física: unidades operativas y su interconexión."
        },
        "glosario": {"terminos": [
            {"termino": "Arquitectura", "definicion": "Atributos visibles al programador."},
            {"termino": "Organización", "definicion": "Unidades operativas y su interconexión."}
        ]},
        "preguntas": {"preguntas": [
            {"pregunta": "¿Qué es la arquitectura de un computador?",
             "respuesta_esperada": "Los atributos visibles al programador.",
             "nivel": "facil"},
            {"pregunta": "¿Cómo se relacionan arquitectura y organización en un mismo procesador?",
             "respuesta_esperada": "La arquitectura se mantiene estable mientras la organización evoluciona con la tecnología.",
             "nivel": "media"},
            {"pregunta": "¿Podría cambiar la organización de un computador sin alterar su arquitectura? Justifica.",
             "respuesta_esperada": "Sí, porque la organización es la implementación física y puede optimizarse sin cambiar lo visible al programador.",
             "nivel": "dificil"}
        ]},
        "actividad": {
            "consigna": "Compara la arquitectura x86 de dos procesadores distintos.",
            "entregable_esperado": "Un cuadro comparativo de una página."
        }
    }

    _agregar_unidad(doc, 1, "Organización vs Arquitectura", recursos_ejemplo, tiene_material=True)
    _agregar_unidad(doc, 2, "Unidad sin material propio", {
        "resumen": {"parrafo_1": "[Generado desde plan docente] La unidad trata sobre memoria caché.",
                    "parrafo_2": "Se explican los niveles L1, L2 y L3 de forma general."},
        "glosario": {"terminos": [{"termino": "Caché", "definicion": "Memoria rápida intermedia."}]},
        "preguntas": {"preguntas": []},
        "actividad": {"consigna": "", "entregable_esperado": ""},
    }, tiene_material=False)

    ruta_logo = "/home/claude/logo_utpl.png"
    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    ruta_prueba = os.path.join(CARPETA_SALIDA, "prueba_exportador.docx")

    if os.path.isfile(ruta_logo):
        with tempfile.TemporaryDirectory() as carpeta_temporal:
            agregar_marca_agua(doc, ruta_logo, "Ing. Juan Pérez Rodríguez", carpeta_temporal)
            doc.save(ruta_prueba)
    else:
        doc.save(ruta_prueba)

    print(f"Documento de prueba generado: {ruta_prueba}")
